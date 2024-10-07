from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import pandas as pd
import os
import zipfile
import json

app = Flask(__name__)

template_file_path = os.path.join('templates', 'product_templates.json')
with open(template_file_path, 'r', encoding='utf-8') as file:
    product_templates = json.load(file)

# Путь к файлу базы штрих-кодов
barcode_file_path = os.path.join('templates', 'База.xlsx')

# Загрузка базы штрих-кодов
barcode_data = pd.read_excel(barcode_file_path, header=None)

# Переименуем столбцы в соответствии с их содержимым
barcode_data.columns = ['Наименование', 'Штрих-код единицы товара', 'Штрих-код упаковки', 'Масса нетто', 'Тара']

# Убираем пустые строки по столбцу "Наименование"
barcode_data.dropna(subset=['Наименование'], inplace=True)

# Путь к файлам для знаков тары
glass_signs_file = os.path.join('templates', 'Знаки_для_тары_Стекло.docx')
doypack_signs_file = os.path.join('templates', 'Знаки_для_тары_Дойпак.txt')
glass_image_file = os.path.join('templates', 'Знаки Тара стекло.Для проверки после Дизайнера.jpg')  # Новый файл

@app.route('/')
def index():
    return render_template('index.html', templates=product_templates)

# Функция для поиска продукта по наименованию и таре
def find_product(name, container_type):
    result = barcode_data[(barcode_data['Наименование'].str.contains(name, case=False, na=False))]
    return result

# Новый маршрут для получения данных шаблона
@app.route('/get_template_data', methods=['POST'])
def get_template_data():
    product_name = request.form.get('product_name')

    # Проверяем, существует ли шаблон с данным именем
    if product_name in product_templates:
        template_data = product_templates[product_name]
        return jsonify(template_data)
    else:
        return jsonify({"error": "Шаблон не найден"}), 404

@app.route('/generate', methods=['POST'])
def generate_label():
    try:
        # Получаем данные из формы
        product_name = request.form.get('product_name')
        container_type = request.form.get('container_type')  # Получаем тип тары
        size = request.form.get('size')
        file_type = request.form.get('file_type')
        doc_type = request.form.get('doc_type')  # Получаем данные из поля ГОСТ/ТУ

        # Проверяем, есть ли шаблон для продукта
        if product_name in product_templates:
            template = product_templates[product_name]
            manufacturer = template['manufacturer']
            location = template['location']
            address = template['address']
            composition = template['composition']
            nutritional_value = template['nutritional_value']
            energi = template['energi']
            storage_conditions = template['storage_conditions']
            net_weight = template['net_weight']
            gost = template['gost']
        else:
            # Если шаблона нет, используем данные, введенные пользователем
            manufacturer = request.form.get('manufacturer')
            location = request.form.get('location')
            address = request.form.get('address')
            composition = request.form.get('composition')
            nutritional_value = request.form.get('nutritional_value')
            energi = request.form.get('energi')
            storage_conditions = request.form.get('storage_conditions')
            net_weight = request.form.get('net_weight')
            gost = request.form.get('gost')

        # Поиск продукта в базе данных по наименованию и таре
        product_data = find_product(product_name, container_type)

        if not product_data.empty:
            product_barcode_unit = product_data.iloc[0]['Штрих-код единицы товара']
            product_barcode_pack = product_data.iloc[0]['Штрих-код упаковки']
            product_mass = product_data.iloc[0]['Масса нетто']
        else:
            product_barcode_unit = 'Не найден'
            product_barcode_pack = 'Не найден'
            product_mass = net_weight

        # Данные для этикетки
        label_data = {
            "Изготовитель": manufacturer,
            "Место нахождения": location,
            "Адрес": address,
            "Продукт": product_name,
            "Состав": composition,
            "Пищевая ценность": nutritional_value,
            "Калоринннннйность": energi,
            "Условия хранения": storage_conditions,
            "Масса нетто": product_mass,
            "Документ": doc_type,  # Используем только текст из поля
            "Штрих-код товара": product_barcode_unit,
            "Штрих-код упаковки": product_barcode_pack
        }

        # Генерация файла этикетки в зависимости от выбранного формата
        if file_type == 'docx':
            doc = Document()
            section = doc.sections[0]

            # Установка размера страницы в зависимости от выбранного формата
            if size == "90 x 60":
                section.page_width = Mm(90)
                section.page_height = Mm(60)
                font_size = Pt(6)  # Установка шрифта 6 pt для размера 90x60
                font_name = "Cambria"  # Шрифт Cambria
            else:
                section.page_width = Mm(58)
                section.page_height = Mm(40)
                font_size = Pt(8)  # Шрифт для других размеров
                font_name = "Arial"  # Основной шрифт для других размеров

            # Установка отступов для более точного расположения текста
            section.top_margin = Mm(5)
            section.bottom_margin = Mm(5)
            section.left_margin = Mm(5)
            section.right_margin = Mm(5)

            # Добавление текста на этикетку
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание по левому краю

            # Устанавливаем межстрочный интервал 1
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Заполняем документ данными
            full_text = ""
            for key, value in label_data.items():
                full_text += f'{key}: {value}\n'

            # Применение шрифта, размера шрифта и интервала
            text_run = paragraph.add_run(full_text)
            text_run.font.size = font_size
            text_run.font.name = font_name

            # Применение стиля к шрифту (Cambria)
            r = text_run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

            # Сохранение документа
            label_filename = f'label_{product_name}.docx'
            label_filepath = os.path.join('generated_labels', label_filename)
            doc.save(label_filepath)

        elif file_type == 'txt':
            label_filename = f'label_{product_name}.txt'
            label_filepath = os.path.join('generated_labels', label_filename)

            # Сохраняем данные в текстовый файл
            with open(label_filepath, 'w', encoding='utf-8') as f:
                for key, value in label_data.items():
                    f.write(f'{key}: {value}\n')

        elif file_type == 'xlsx':
            label_filename = f'label_{product_name}.xlsx'
            label_filepath = os.path.join('generated_labels', label_filename)

            # Используем pandas для записи данных в Excel
            df = pd.DataFrame([label_data])
            df.to_excel(label_filepath, index=False)

        else:
            return f"Неподдерживаемый формат файла: {file_type}", 400

        # Определяем файл со знаками для выбранной тары
        if container_type == 'glass':
            signs_file = glass_signs_file
            # Добавляем файл изображения стекла, если выбрано "Стекло"
            image_file = glass_image_file
        elif container_type == 'doypack':
            signs_file = doypack_signs_file
            image_file = None  # Для дойпака не добавляем изображение

        # Создаем ZIP архив для скачивания обоих файлов
        zip_filename = f'label_and_signs_{product_name}.zip'
        zip_filepath = os.path.join('generated_labels', zip_filename)

        with zipfile.ZipFile(zip_filepath, 'w') as zipf:
            # Добавляем файл этикетки
            zipf.write(label_filepath, arcname=label_filename)
            # Добавляем файл знаков
            zipf.write(signs_file, arcname=os.path.basename(signs_file))
            # Если выбрана тара "Стекло", добавляем изображение
            if image_file:
                zipf.write(image_file, arcname=os.path.basename(image_file))

        # Отправляем ZIP файл пользователю
        return send_file(zip_filepath, as_attachment=True)

    except Exception as e:
        return f"Произошла ошибка: {e}"

    # Создаем папку для сохранения этикеток, если её нет
    if not os.path.exists('generated_labels'):
        os.makedirs('generated_labels')


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=80, debug=True)
