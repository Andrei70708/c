from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import os
import zipfile
import json

app = Flask(__name__)

# Путь к файлу шаблонов
template_file_path = os.path.join('templates', 'product_templates.json')

# Загрузка шаблонов
with open(template_file_path, 'r', encoding='utf-8') as file:
    product_templates = json.load(file)

# Путь к файлу базы штрих-кодов
barcode_file_path = os.path.join('templates', 'БАЗА.xlsx')

# Загрузка базы штрих-кодов
barcode_data = pd.read_excel(barcode_file_path, header=None)

# Переименуем столбцы в соответствии с их содержимым
barcode_data.columns = ['Наименование', 'Штрих-код единицы товара', 'Штрих-код упаковки', 'Масса нетто', 'Тара']

# Убираем пустые строки по столбцу "Наименование"
barcode_data.dropna(subset=['Наименование'], inplace=True)

# Путь к файлам для знаков тары
glass_signs_file = os.path.join('templates', 'Знаки_для_тары_Стекло.txt')
doypack_signs_file = os.path.join('templates', 'Знаки_для_тары_Дойпак.txt')


@app.route('/')
def index():
    return render_template('index.html', templates=product_templates)

# Функция для поиска продукта по наименованию и таре
def find_product(name, container_type):
    result = barcode_data[(barcode_data['Наименование'].str.contains(name, case=False, na=False)) &
                          (barcode_data['Тара'].str.contains(container_type, case=False, na=False))]
    return result

# Новый маршрут для получения данных шаблона
@app.route('/get_template_data', methods=['POST'])
def get_template_data():
    product_name = request.form.get('product_name')

    # Проверяем, существует ли шаблон с данным именем
    if product_name in product_templates:
        template_data = product_templates[product_name]
        return jsonify(template_data)
    else:
        return jsonify({"error": "Шаблон не найден"}), 404

@app.route('/generate', methods=['POST'])
def generate_label():
    try:
        # Получаем данные из формы
        product_name = request.form.get('product_name')
        container_type = request.form.get('container_type')  # Получаем тип тары
        size = request.form.get('size')
        file_type = request.form.get('file_type')

        # Проверяем, есть ли шаблон для продукта
        if product_name in product_templates:
            template = product_templates[product_name]
            manufacturer = template['manufacturer']
            location = template['location']
            address = template['address']
            composition = template['composition']
            nutritional_value = template['nutritional_value']
            energi = template['energi']
            storage_conditions = template['storage_conditions']
            net_weight = template['net_weight']
            gost = template['gost']
        else:
            # Если шаблона нет, используем данные, введенные пользователем
            manufacturer = request.form.get('manufacturer')
            location = request.form.get('location')
            address = request.form.get('address')
            composition = request.form.get('composition')
            nutritional_value = request.form.get('nutritional_value')
            energi = request.form.get('energi')
            storage_conditions = request.form.get('storage_conditions')
            net_weight = request.form.get('net_weight')
            gost = request.form.get('gost')

        # Поиск продукта в базе данных по наименованию и таре
        product_data = find_product(product_name, container_type)

        if not product_data.empty:
            product_barcode_unit = product_data.iloc[0]['Штрих-код единицы товара']
            product_barcode_pack = product_data.iloc[0]['Штрих-код упаковки']
            product_mass = product_data.iloc[0]['Масса нетто']
        else:
            product_barcode_unit = 'Не найден'
            product_barcode_pack = 'Не найден'
            product_mass = net_weight

        # Данные для этикетки
        label_data = {
            "Изготовитель": manufacturer,
            "Место нахождения": location,
            "Адрес": address,
            "Продукт": product_name,
            "Состав": composition,
            "Пищевая ценность": nutritional_value,
            "Калорийность": energi,
            "Условия хранения": storage_conditions,
            "Масса нетто": product_mass,
            "ГОСТ": gost,
            "Штрих-код товара": product_barcode_unit,
            "Штрих-код упаковки": product_barcode_pack
        }

        # Генерация файла этикетки в зависимости от выбранного формата
        if file_type == 'docx':
            doc = Document()
            section = doc.sections[0]

            # Установка размера страницы
            if size == "90 x 60":
                section.page_width = Mm(90)
                section.page_height = Mm(60)
            elif size == "58 x 40":
                section.page_width = Mm(58)
                section.page_height = Mm(40)

            # Добавление текста на этикетку
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Заполняем документ данными
            for key, value in label_data.items():
                paragraph.add_run(f'{key}: {value}\n').font.size = Pt(10)

            # Сохранение документа
            label_filename = f'label_{product_name}.docx'
            label_filepath = os.path.join('generated_labels', label_filename)
            doc.save(label_filepath)

        elif file_type == 'txt':
            label_filename = f'label_{product_name}.txt'
            label_filepath = os.path.join('generated_labels', label_filename)

            # Сохраняем данные в текстовый файл
            with open(label_filepath, 'w', encoding='utf-8') as f:
                for key, value in label_data.items():
                    f.write(f'{key}: {value}\n')

        elif file_type == 'xlsx':
            label_filename = f'label_{product_name}.xlsx'
            label_filepath = os.path.join('generated_labels', label_filename)

            # Используем pandas для записи данных в Excel
            df = pd.DataFrame([label_data])
            df.to_excel(label_filepath, index=False)

        else:
            return f"Неподдерживаемый формат файла: {file_type}", 400

        # Определяем файл со знаками для выбранной тары
        if container_type == 'glass':
            signs_file = glass_signs_file
        elif container_type == 'doypack':
            signs_file = doypack_signs_file

        # Создаем ZIP архив для скачивания обоих файлов
        zip_filename = f'label_and_signs_{product_name}.zip'
        zip_filepath = os.path.join('generated_labels', zip_filename)

        with zipfile.ZipFile(zip_filepath, 'w') as zipf:
            # Добавляем файл этикетки
            zipf.write(label_filepath, arcname=label_filename)
            # Добавляем файл знаков
            zipf.write(signs_file, arcname=os.path.basename(signs_file))

        # Отправляем ZIP файл пользователю
        return send_file(zip_filepath, as_attachment=True)

    except Exception as e:
        return f"Произошла ошибка: {e}"


if __name__ == '__main__':
    # Создаем папку для сохранения этикеток, если её нет
    if not os.path.exists('generated_labels'):
        os.makedirs('generated_labels')

    app.run(debug=True)
