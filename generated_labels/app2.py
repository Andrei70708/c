from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import os

app = Flask(__name__)

# Шаблоны данных для различных продуктов
templates = {
    "Кетчуп": {
        "manufacturer": "«Буздякский консервный комбинат»-филиал АО«Пищепром»",
        "location": "450022, Россия, Республика Башкортостан,  г. Уфа, ул. Менделеева, д. 153",
        "address": "452710, Россия, Республика Башкортостан, Буздякский район, с. Буздяк, ул. Гагарина, д. 45.",
        "product_name": "Кетчуп Острый первой категории",
        "composition": "вода, томатная паста...",
        "may_contain": "Может содержать следы глютена.",
        "nutritional_value": "белки —3,3 г, жиры — 0,3 г, углеводы – 13,6 г.",
        "energi": "70 ккал/293 кДж.",
        "storage_conditions": "Хранить при температуре от 0°С до +25°С.",
        "net_weight": "260 г.",
        "gost": "ГОСТ 32063"
    }
    # Добавь другие шаблоны здесь...
}

@app.route('/')
def index():
    return render_template('index.html', templates=templates)

@app.route('/generate', methods=['POST'])
def generate_label():
    try:
        # Получаем данные из формы
        manufacturer = request.form.get('manufacturer')
        location = request.form.get('location')
        address = request.form.get('address')
        product_name = request.form.get('product_name')
        composition = request.form.get('composition')
        may_contain = request.form.get('may_contain')
        nutritional_value = request.form.get('nutritional_value')
        energi = request.form.get('energi')
        storage_conditions = request.form.get('storage_conditions')
        net_weight = request.form.get('net_weight')
        gost = request.form.get('gost')
        size = request.form.get('size')
        file_type = request.form.get('file_type')

        # Данные для этикетки
        label_data = {
            "Изготовитель": manufacturer,
            "Место нахождения": location,
            "Адрес": address,
            "Продукт": product_name,
            "Состав": composition,
            "Может содержать": may_contain,
            "Пищевая ценность": nutritional_value,
            "Калорийность": energi,
            "Условия хранения": storage_conditions,
            "Масса нетто": net_weight,
            "ГОСТ": gost
        }

        # Генерация этикетки в формате Word (.docx)
        if file_type == 'docx':
            doc = Document()
            section = doc.sections[0]

            # Установка размера страницы
            if size == "90 x 60":
                section.page_width = Mm(90)
                section.page_height = Mm(60)
            elif size == "58 x 40":
                section.page_width = Mm(58)
                section.page_height = Mm(40)

            # Добавление текста на этикетку
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Заполняем документ данными
            for key, value in label_data.items():
                paragraph.add_run(f'{key}: {value}\n').font.size = Pt(10)

            # Сохранение документа
            filename = f'label_{product_name}.docx'
            filepath = os.path.join('generated_labels', filename)
            doc.save(filepath)

            # Возвращаем файл пользователю для скачивания
            return send_file(filepath, as_attachment=True)

        # Генерация этикетки в формате TXT
        elif file_type == 'txt':
            filename = f'label_{product_name}.txt'
            filepath = os.path.join('generated_labels', filename)

            # Сохраняем данные в текстовый файл
            with open(filepath, 'w', encoding='utf-8') as f:
                for key, value in label_data.items():
                    f.write(f'{key}: {value}\n')

            return send_file(filepath, as_attachment=True)

        # Генерация этикетки в формате Excel (.xlsx)
        elif file_type == 'xlsx':
            filename = f'label_{product_name}.xlsx'
            filepath = os.path.join('generated_labels', filename)

            # Используем pandas для записи данных в Excel
            df = pd.DataFrame([label_data])
            df.to_excel(filepath, index=False)

            return send_file(filepath, as_attachment=True)

        else:
            return f"Неподдерживаемый формат файла: {file_type}", 400

    except Exception as e:
        return f"Произошла ошибка: {e}"

if __name__ == '__main__':
    # Создаем папку для сохранения этикеток, если её нет
    if not os.path.exists('generated_labels'):
        os.makedirs('generated_labels')

    app.run(debug=True)
